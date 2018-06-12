# Text in DOCX files
# pip install python-docx

import docx

document = docx.Document('example.docx')

print(len(document.paragraphs))

# Getting all the text from the first paragraph
print(document.paragraphs[0].text)

print(document.paragraphs[1].text)

# Getting a specific run from a specific paragraph
print(document.paragraphs[0].runs[19].text)

# Getting all text from a DOCX file #

def getText(filename):
    doc = docx.Document(filename)
    text = []
    for p in doc.paragraphs:
        text.append(p.text)

    return '\n'.join(text)

print("==================================================")

print(getText('nf.docx'))
