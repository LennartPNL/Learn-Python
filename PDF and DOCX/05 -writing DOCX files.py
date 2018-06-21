import docx

d = docx.Document()

# Adds the text given as parameter to a new paragraph
d.add_paragraph('Hello there!')

newParagraph = d.add_paragraph('Here\'s another paragraph for you')

# Adds the text after the given paragraph
newParagraph.add_run(', even with some added text to it.')

# Save the file
d.save('write.docx')

# Adding Headings #

# Adds the biggest header possible to the document
d.add_heading('This is a huge-ass header mate', 0)

# A smaller header
d.add_heading('But you can go smaller if you want', 1)

d.save('write.docx')

# Pictures #

# Adding a picture and giving it a size
d.add_picture('dog.jpg', width=docx.shared.Cm(10), height=docx.shared.Cm(10))

d.save('write.docx')
